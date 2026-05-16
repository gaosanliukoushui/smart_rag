"""Unit tests for parsers."""

import pytest
from pathlib import Path

from app.parsers.text_parser import TextParser
from app.parsers.markdown_parser import MarkdownParser
from app.parsers import get_parser, get_supported_extensions


class TestTextParser:
    """Tests for TextParser."""

    def test_parse_basic(self, temp_dir, sample_text):
        """Test basic text parsing."""
        file_path = temp_dir / "test.txt"
        file_path.write_text(sample_text, encoding="utf-8")

        parser = TextParser()
        result = parser.parse(file_path)

        assert "测试文本" in result
        assert "第一段落" in result

    def test_get_metadata(self, temp_dir, sample_text):
        """Test metadata extraction."""
        file_path = temp_dir / "test.txt"
        file_path.write_text(sample_text, encoding="utf-8")

        parser = TextParser()
        metadata = parser.get_metadata(file_path)

        assert metadata["title"] == "test"
        assert "lines" in metadata
        assert "word_count" in metadata


class TestMarkdownParser:
    """Tests for MarkdownParser."""

    def test_parse_basic(self, temp_dir):
        """Test basic markdown parsing."""
        content = """# Title

This is a paragraph.

## Section

Another paragraph.
"""
        file_path = temp_dir / "test.md"
        file_path.write_text(content, encoding="utf-8")

        parser = MarkdownParser()
        result = parser.parse(file_path)

        assert "Title" in result
        assert "paragraph" in result
        assert "#" not in result

    def test_get_metadata(self, temp_dir):
        """Test markdown metadata extraction."""
        content = """# Document Title

Some content.
"""
        file_path = temp_dir / "test.md"
        file_path.write_text(content, encoding="utf-8")

        parser = MarkdownParser()
        metadata = parser.get_metadata(file_path)

        assert metadata["title"] == "Document Title"
        assert metadata["lines"] > 0


class TestPDFParser:
    """Tests for PDFParser."""

    @pytest.fixture
    def pdf_parser(self):
        pytest.importorskip("pypdf", reason="pypdf not installed")
        from app.parsers.pdf_parser import PDFParser, PYPDF_AVAILABLE
        if not PYPDF_AVAILABLE:
            pytest.skip("pypdf not available")
        return PDFParser()

    def test_parse_basic(self, temp_dir, pdf_parser):
        """Test PDF parsing extracts text from pages."""
        try:
            import pypdf
        except ImportError:
            pytest.skip("pypdf not installed")

        pdf_path = temp_dir / "sample.pdf"
        writer = pypdf.PdfWriter()
        writer.add_blank_page(width=200, height=200)
        writer.write(pdf_path.open("wb"))

        result = pdf_parser.parse(pdf_path)
        assert isinstance(result, str)

    def test_get_metadata(self, temp_dir, pdf_parser):
        """Test PDF metadata extraction."""
        try:
            import pypdf
        except ImportError:
            pytest.skip("pypdf not installed")

        pdf_path = temp_dir / "sample.pdf"
        writer = pypdf.PdfWriter()
        writer.add_blank_page(width=200, height=200)
        writer.write(pdf_path.open("wb"))

        metadata = pdf_parser.get_metadata(pdf_path)
        assert "pages" in metadata
        assert metadata["pages"] >= 1


class TestWordParser:
    """Tests for WordParser."""

    @pytest.fixture
    def word_parser(self):
        pytest.importorskip("docx", reason="python-docx not installed")
        from app.parsers.word_parser import WordParser, DOCX_AVAILABLE
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx not available")
        return WordParser()

    def test_parse_basic(self, temp_dir, word_parser):
        """Test Word document parsing extracts paragraph text."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        docx_path = temp_dir / "sample.docx"
        doc = DocxDocument()
        doc.add_paragraph("First paragraph text.")
        doc.add_paragraph("Second paragraph text.")
        doc.save(str(docx_path))

        result = word_parser.parse(docx_path)
        assert "First paragraph text" in result
        assert "Second paragraph text" in result

    def test_parse_with_table(self, temp_dir, word_parser):
        """Test Word document table content is extracted."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        docx_path = temp_dir / "with_table.docx"
        doc = DocxDocument()
        doc.add_paragraph("Before table.")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"
        doc.save(str(docx_path))

        result = word_parser.parse(docx_path)
        assert "A1" in result
        assert "B1" in result
        assert "A2" in result
        assert "B2" in result

    def test_get_metadata(self, temp_dir, word_parser):
        """Test Word metadata extraction."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        docx_path = temp_dir / "meta.docx"
        doc = DocxDocument()
        doc.add_paragraph("Para 1")
        doc.add_paragraph("Para 2")
        table = doc.add_table(rows=1, cols=1)
        doc.save(str(docx_path))

        metadata = word_parser.get_metadata(docx_path)
        assert metadata["paragraphs"] >= 2
        assert metadata["tables"] >= 1


class TestParserFactory:
    """Tests for parser factory functions."""

    def test_get_parser_txt(self, temp_dir):
        """Test get_parser returns TextParser for .txt files."""
        from app.parsers.text_parser import TextParser
        parser = get_parser(temp_dir / "file.txt")
        assert isinstance(parser, TextParser)

    def test_get_parser_md(self, temp_dir):
        """Test get_parser returns MarkdownParser for .md files."""
        from app.parsers.markdown_parser import MarkdownParser
        parser = get_parser(temp_dir / "file.md")
        assert isinstance(parser, MarkdownParser)

    def test_get_parser_markdown(self, temp_dir):
        """Test get_parser returns MarkdownParser for .markdown files."""
        from app.parsers.markdown_parser import MarkdownParser
        parser = get_parser(temp_dir / "file.markdown")
        assert isinstance(parser, MarkdownParser)

    def test_get_parser_unsupported(self, temp_dir):
        """Test get_parser raises ValueError for unsupported file types."""
        with pytest.raises(ValueError) as exc_info:
            get_parser(temp_dir / "file.xyz")
        assert "Unsupported file type" in str(exc_info.value)
        assert ".xyz" in str(exc_info.value)

    def test_get_supported_extensions(self):
        """Test get_supported_extensions returns all extensions."""
        exts = get_supported_extensions()
        assert ".txt" in exts
        assert ".md" in exts
        assert ".pdf" in exts
        assert ".docx" in exts

    def test_case_insensitive(self, temp_dir):
        """Test get_parser is case-insensitive for extensions."""
        from app.parsers.markdown_parser import MarkdownParser
        from app.parsers.text_parser import TextParser
        parser_upper = get_parser(temp_dir / "file.TXT")
        parser_lower = get_parser(temp_dir / "file.txt")
        assert isinstance(parser_upper, TextParser)
        assert isinstance(parser_lower, TextParser)

    def test_get_parser_docx(self, temp_dir):
        """Test get_parser returns WordParser for .docx files."""
        from app.parsers.word_parser import WordParser
        parser = get_parser(temp_dir / "file.docx")
        assert isinstance(parser, WordParser)

    def test_get_parser_pdf(self, temp_dir):
        """Test get_parser returns PDFParser for .pdf files."""
        from app.parsers.pdf_parser import PDFParser
        parser = get_parser(temp_dir / "file.pdf")
        assert isinstance(parser, PDFParser)
