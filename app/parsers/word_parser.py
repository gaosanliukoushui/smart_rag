"""Word document parser."""

from pathlib import Path
from typing import Dict, Any

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.parsers.base import BaseParser


class WordParser(BaseParser):
    """Parser for Word documents (.docx)."""

    SUPPORTED_EXTENSIONS = [".docx"]

    def parse(self, file_path: str | Path) -> str:
        """Extract text content from a Word document."""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for Word parsing. Install with: pip install python-docx"
            )

        file_path = Path(file_path)
        text_parts = []

        doc = DocxDocument(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    def get_metadata(self, file_path: str | Path) -> Dict[str, Any]:
        """Extract metadata from a Word document."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word parsing")

        file_path = Path(file_path)
        metadata = {
            "paragraphs": 0,
            "tables": 0,
            "title": file_path.stem,
        }

        doc = DocxDocument(file_path)
        metadata["paragraphs"] = len([p for p in doc.paragraphs if p.text.strip()])
        metadata["tables"] = len(doc.tables)

        core_props = doc.core_properties
        if core_props.title:
            metadata["title"] = core_props.title

        return metadata
